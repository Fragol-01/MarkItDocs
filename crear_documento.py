from __future__ import annotations

import argparse
import base64
import glob
import json
import logging
import re
import time
from dataclasses import dataclass
from io import BytesIO
from pathlib import Path
from urllib.error import URLError
from urllib.parse import unquote, urlparse
from urllib.request import Request, urlopen

import markdown as markdown_lib
from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from lxml import html as lxml_html

try:
    import tomllib
except ModuleNotFoundError:  # pragma: no cover - Python <3.11 fallback
    tomllib = None

logger = logging.getLogger("crear_documento")
if not logger.handlers:
    _handler = logging.StreamHandler()
    _handler.setFormatter(logging.Formatter("%(levelname)s: %(message)s"))
    logger.addHandler(_handler)
    logger.setLevel(logging.INFO)

MARKDOWN_EXTENSIONS = {".md", ".markdown"}
HTML_EXTENSIONS = {".html", ".htm"}
SUPPORTED_EXTENSIONS = MARKDOWN_EXTENSIONS | HTML_EXTENSIONS

DEFAULT_THEME = {
    "body_font": "Calibri",
    "code_font": "Consolas",
    "title_color": [31, 78, 121],
    "link_color": [0, 102, 204],
    "quote_fill": "F8F9FA",
    "table_head_fill": "D9EAF7",
    "code_fill": "F6F8FA",
    "hr_color": [180, 190, 205],
}

BODY_FONT = DEFAULT_THEME["body_font"]
CODE_FONT = DEFAULT_THEME["code_font"]
TITLE_COLOR = RGBColor(*DEFAULT_THEME["title_color"])
LINK_COLOR = RGBColor(*DEFAULT_THEME["link_color"])
QUOTE_FILL = DEFAULT_THEME["quote_fill"]
TABLE_HEAD_FILL = DEFAULT_THEME["table_head_fill"]
CODE_FILL = DEFAULT_THEME["code_fill"]
HR_COLOR = RGBColor(*DEFAULT_THEME["hr_color"])

PAGE_BREAK_PATTERN = re.compile(r"(?im)^(?:\\pagebreak|<!--\s*pagebreak\s*-->)\s*$")
MAX_IMAGE_DOWNLOAD_RETRIES = 3
RETRY_BACKOFF_SECONDS = 1.5

HEADING_TAGS = {f"h{i}" for i in range(1, 7)}
BLOCK_TAGS = {
    "article",
    "aside",
    "blockquote",
    "div",
    "figure",
    "figcaption",
    "li",
    "nav",
    "ol",
    "p",
    "pre",
    "section",
    "table",
    "tbody",
    "thead",
    "tfoot",
    "tr",
    "ul",
}
TOC_PATTERN = re.compile(r"(?im)^(?:\[toc\]|\[\[toc\]\]|<!--\s*toc\s*-->)\s*$")


@dataclass(frozen=True)
class InlineStyle:
    bold: bool = False
    italic: bool = False
    code: bool = False
    strike: bool = False
    underline: bool = False
    color: RGBColor | None = None


def slugify(text: str) -> str:
    value = re.sub(r"[^0-9a-zA-Z]+", "_", text.strip().lower())
    value = re.sub(r"_+", "_", value).strip("_")
    return value or "section"


def unique_name(base: str, used: set[str]) -> str:
    candidate = base
    index = 2
    while candidate in used:
        candidate = f"{base}_{index}"
        index += 1
    used.add(candidate)
    return candidate


def color_to_hex(color: RGBColor) -> str:
    return f"{color[0]:02X}{color[1]:02X}{color[2]:02X}"


def load_theme(theme_path: str | Path | None) -> dict:
    """Carga un tema de estilos.

    Acepta:
    - ``None`` → tema por defecto
    - Nombre de tema integrado (``professional``, ``academico``, …) → usa la
      metadata JSON/YAML del subpaquete markitpdf, así el mismo tema viste
      tanto el PDF como el Word.
    - Ruta a un archivo ``.json``/``.toml``/``.yaml``/``.yml`` con overrides.
    """
    if theme_path is None:
        return dict(DEFAULT_THEME)

    theme_path = Path(theme_path)
    if not theme_path.exists():
        # ¿Es el nombre de un tema integrado en vez de una ruta?
        name = str(theme_path)
        if "/" not in name and "\\" not in name and not theme_path.suffix:
            from markitpdf.converter import available_themes, get_theme_metadata

            if name in available_themes():
                theme = dict(DEFAULT_THEME)
                theme.update(get_theme_metadata(name).to_docx_theme())
                return theme
            raise ValueError(
                f"Tema '{name}' no reconocido. Integrados: {', '.join(available_themes())}, "
                "o pasa la ruta de un archivo .json/.toml/.yaml"
            )
        raise FileNotFoundError(f"No existe el archivo de tema: {theme_path}")

    suffix = theme_path.suffix.lower()
    if suffix == ".json":
        overrides = json.loads(theme_path.read_text(encoding="utf-8"))
    elif suffix == ".toml":
        if tomllib is None:
            raise RuntimeError("Se requiere Python 3.11+ para leer temas .toml")
        overrides = tomllib.loads(theme_path.read_text(encoding="utf-8"))
    elif suffix in {".yaml", ".yml"}:
        try:
            import yaml  # type: ignore
        except ImportError as exc:
            raise RuntimeError(
                f"Para cargar {theme_path.name} instala PyYAML: pip install pyyaml"
            ) from exc
        overrides = yaml.safe_load(theme_path.read_text(encoding="utf-8"))
    else:
        raise ValueError(
            f"Formato de tema no soportado: {theme_path.suffix} (usa .json/.toml/.yaml/.yml)"
        )

    if not isinstance(overrides, dict):
        raise ValueError(f"El tema {theme_path.name} debe contener un objeto/mapping en la raíz.")

    theme = dict(DEFAULT_THEME)
    theme.update(overrides)
    return theme


def configure_document(document: Document, theme: dict) -> None:
    section = document.sections[0]
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)

    title_color = RGBColor(*theme["title_color"])

    normal = document.styles["Normal"]
    normal.font.name = theme["body_font"]
    normal.font.size = Pt(10.5)

    for style_name, size in {
        "Heading 1": 20,
        "Heading 2": 16,
        "Heading 3": 13,
        "Heading 4": 11.5,
        "Heading 5": 10.5,
        "Heading 6": 10,
    }.items():
        style = document.styles[style_name]
        style.font.name = theme["body_font"]
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = title_color


def add_paragraph_shading(paragraph, fill: str) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    p_pr.append(shd)


def add_paragraph_border(paragraph, color: str = "D0D7DE") -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), color)
    p_bdr.append(bottom)
    p_pr.append(p_bdr)


def add_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def add_hyperlink(
    paragraph,
    text: str,
    target: str,
    *,
    internal: bool = False,
    style: InlineStyle | None = None,
    link_color: RGBColor | None = None,
    code_font: str | None = None,
) -> None:
    style = style or InlineStyle()
    link_color = link_color or LINK_COLOR
    code_font = code_font or CODE_FONT
    hyperlink = OxmlElement("w:hyperlink")
    if internal:
        hyperlink.set(qn("w:anchor"), target)
    else:
        rel_id = paragraph.part.relate_to(
            target,
            "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
            is_external=True,
        )
        hyperlink.set(qn("r:id"), rel_id)

    run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), color_to_hex(style.color or link_color))
    r_pr.append(color)

    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    r_pr.append(underline)

    if style.bold:
        r_pr.append(OxmlElement("w:b"))
    if style.italic:
        r_pr.append(OxmlElement("w:i"))
    if style.strike:
        r_pr.append(OxmlElement("w:strike"))
    if style.code:
        fonts = OxmlElement("w:rFonts")
        fonts.set(qn("w:ascii"), code_font)
        fonts.set(qn("w:hAnsi"), code_font)
        fonts.set(qn("w:cs"), code_font)
        r_pr.append(fonts)

    run.append(r_pr)
    text_node = OxmlElement("w:t")
    text_node.set(qn("xml:space"), "preserve")
    text_node.text = text
    run.append(text_node)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def add_bookmark(paragraph, name: str, bookmark_id: int) -> None:
    start = OxmlElement("w:bookmarkStart")
    start.set(qn("w:id"), str(bookmark_id))
    start.set(qn("w:name"), name)
    end = OxmlElement("w:bookmarkEnd")
    end.set(qn("w:id"), str(bookmark_id))
    paragraph._p.insert(0, start)
    paragraph._p.append(end)


def add_toc_field(paragraph, start_level: int = 1, end_level: int = 3) -> None:
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = f'TOC \\o "{start_level}-{end_level}" \\h \\z \\u'
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    placeholder = OxmlElement("w:t")
    placeholder.text = "Actualiza la tabla de contenido en Word."
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.append(begin)
    run._r.append(instr)
    run._r.append(separate)
    run._r.append(placeholder)
    run._r.append(end)


# Una fila de tabla pipe y su separador de cabecera (| :--- | --- |)
_TABLE_ROW_RE = re.compile(r"^\s{0,3}\|.*\|\s*$")
_TABLE_SEP_RE = re.compile(r"^\s{0,3}\|?(\s*:?-+:?\s*\|)+\s*:?-*:?\s*\|?\s*$")


def normalize_table_blank_lines(markdown_text: str) -> str:
    """Inserta la línea en blanco que python-markdown exige antes de una tabla.

    Mucha gente escribe la tabla pegada al párrafo anterior (GitHub y VS Code
    la renderizan igual), pero la extensión 'tables' la ignora en silencio.
    Detectamos "fila | separador" y garantizamos la línea en blanco previa.
    (Duplicado a propósito en markitpdf/converter.py: los dos motores son
    independientes y no queremos acoplarlos por 15 líneas.)
    """
    lines = markdown_text.splitlines()
    out: list[str] = []
    for i, line in enumerate(lines):
        if (
            _TABLE_ROW_RE.match(line)
            and i + 1 < len(lines)
            and _TABLE_SEP_RE.match(lines[i + 1])
            and out
            and out[-1].strip()
            and not _TABLE_ROW_RE.match(out[-1])
        ):
            out.append("")
        out.append(line)
    result = "\n".join(out)
    if markdown_text.endswith("\n") and not result.endswith("\n"):
        result += "\n"
    return result


def markdown_to_html(markdown_text: str) -> str:
    return markdown_lib.markdown(
        normalize_table_blank_lines(markdown_text),
        extensions=["extra", "tables", "fenced_code", "attr_list", "sane_lists", "toc"],
        output_format="html5",
    )


def is_remote_source(src: str) -> bool:
    return urlparse(src).scheme.lower() in {"http", "https"}


def read_image_bytes(src: str, source_dir: Path) -> tuple[BytesIO | None, str | None]:
    src = src.strip()
    if not src:
        return None, None

    if src.startswith("data:"):
        match = re.match(r"data:[^;]+;base64,(.+)", src, re.IGNORECASE | re.DOTALL)
        if not match:
            return None, None
        return BytesIO(base64.b64decode(match.group(1))), None

    if is_remote_source(src):
        request = Request(src, headers={"User-Agent": "Mozilla/5.0"})
        last_error: Exception | None = None
        for attempt in range(1, MAX_IMAGE_DOWNLOAD_RETRIES + 1):
            try:
                with urlopen(request, timeout=20) as response:
                    return BytesIO(response.read()), None
            except (URLError, TimeoutError, OSError) as exc:
                last_error = exc
                logger.warning(
                    "Descarga de imagen falló (intento %d/%d) %s: %s",
                    attempt, MAX_IMAGE_DOWNLOAD_RETRIES, src, exc,
                )
                if attempt < MAX_IMAGE_DOWNLOAD_RETRIES:
                    time.sleep(RETRY_BACKOFF_SECONDS * attempt)
        return None, f"{src} (descarga falló tras {MAX_IMAGE_DOWNLOAD_RETRIES} intentos: {last_error})"

    image_path = (source_dir / unquote(src)).resolve()
    if not image_path.exists():
        return None, str(image_path)
    return BytesIO(image_path.read_bytes()), str(image_path)


def apply_run_style(run, style: InlineStyle, body_font: str | None = None, code_font: str | None = None) -> None:
    body_font = body_font or BODY_FONT
    code_font = code_font or CODE_FONT
    run.font.name = code_font if style.code else body_font
    run.font.size = Pt(9.5 if style.code else 10.5)
    run.bold = style.bold
    run.italic = style.italic
    run.font.strike = style.strike
    run.underline = style.underline
    if style.color is not None:
        run.font.color.rgb = style.color


def add_text(
    paragraph,
    text: str,
    style: InlineStyle = InlineStyle(),
    body_font: str | None = None,
    code_font: str | None = None,
) -> None:
    if not text:
        return
    run = paragraph.add_run(text)
    apply_run_style(run, style, body_font=body_font, code_font=code_font)


def text_content(node) -> str:
    return "".join(node.itertext()).strip()


class MarkdownToDocxConverter:
    def __init__(self, source_path: Path, output_path: Path, theme: dict | None = None) -> None:
        self.source_path = source_path
        self.output_path = output_path
        self.source_dir = source_path.parent
        self.theme = theme or DEFAULT_THEME
        self.title_color = RGBColor(*self.theme["title_color"])
        self.link_color = RGBColor(*self.theme["link_color"])
        self.body_font = self.theme["body_font"]
        self.code_font = self.theme["code_font"]
        self.document = Document()
        configure_document(self.document, self.theme)
        self.available_width = (
            self.document.sections[0].page_width
            - self.document.sections[0].left_margin
            - self.document.sections[0].right_margin
        )
        self.bookmarks: set[str] = set()
        self.anchor_map: dict[str, str] = {}
        self.bookmark_id = 1

    def convert(self, markdown_text: str) -> Document:
        return self.convert_markdown(markdown_text)

    def convert_html(self, html_text: str) -> Document:
        """Convierte contenido HTML (documento completo o fragmento) a .docx."""
        # Si trae <body>, nos quedamos con su interior; si no, lo usamos tal cual.
        body_match = re.search(r"<body[^>]*>(.*)</body>", html_text, re.IGNORECASE | re.DOTALL)
        body = body_match.group(1) if body_match else html_text
        # Dividimos por hr.__pagebreak__ para soportar saltos de página manuales
        # (insertamos esta clase en HTML a través de un comentario o un <hr>).
        parts = re.split(r'<hr[^>]*class="__pagebreak__"[^>]*>', body, flags=re.IGNORECASE)
        combined = '<hr class="__pagebreak__">'.join(parts)
        root = lxml_html.fromstring(f"<div>{combined}</div>")
        self._build_anchor_map(root)
        self._render_toc_if_requested_html(root, html_text)
        self._render_container(root)
        return self.document

    def convert_markdown(self, markdown_text: str) -> Document:
        segments = PAGE_BREAK_PATTERN.split(markdown_text)
        html_segments = [markdown_to_html(segment) for segment in segments]
        combined_html = "<hr class=\"__pagebreak__\">".join(html_segments)
        root = lxml_html.fromstring(f"<div>{combined_html}</div>")
        self._build_anchor_map(root)
        self._render_toc_if_requested(markdown_text)
        self._render_container(root)
        return self.document

    def _render_toc_if_requested_html(self, root, html_text: str) -> None:
        if not TOC_PATTERN.search(html_text):
            return
        title = self.document.add_paragraph()
        title.style = "Heading 1"
        title.add_run("Tabla de contenido")
        title.runs[0].font.color.rgb = self.title_color
        toc = self.document.add_paragraph()
        add_toc_field(toc)

    def _build_anchor_map(self, root) -> None:
        for node in root.iterdescendants():
            if node.tag not in HEADING_TAGS:
                continue
            heading_text = text_content(node)
            anchor = node.get("id") or slugify(heading_text)
            bookmark = unique_name(slugify(anchor), self.bookmarks)
            self.anchor_map[anchor] = bookmark
            self.anchor_map[slugify(anchor)] = bookmark
            self.anchor_map[slugify(heading_text)] = bookmark

    def _bookmark_for(self, anchor: str) -> str | None:
        clean = anchor.lstrip("#")
        return self.anchor_map.get(clean) or self.anchor_map.get(slugify(clean))

    def _next_bookmark_id(self) -> int:
        value = self.bookmark_id
        self.bookmark_id += 1
        return value

    def _add_text(self, paragraph, text: str, style: InlineStyle = InlineStyle()) -> None:
        add_text(paragraph, text, style, body_font=self.body_font, code_font=self.code_font)

    def _add_hyperlink(self, paragraph, text: str, target: str, *, internal: bool = False, style: InlineStyle | None = None) -> None:
        add_hyperlink(
            paragraph, text, target,
            internal=internal, style=style,
            link_color=self.link_color, code_font=self.code_font,
        )

    def _render_toc_if_requested(self, markdown_text: str) -> None:
        if not TOC_PATTERN.search(markdown_text):
            return
        title = self.document.add_paragraph()
        title.style = "Heading 1"
        title.add_run("Tabla de contenido")
        title.runs[0].font.color.rgb = self.title_color
        toc = self.document.add_paragraph()
        add_toc_field(toc)

    def _render_container(self, root) -> None:
        for child in root:
            self._render_block(child)

    def _render_block(self, node) -> None:
        tag = getattr(node, "tag", None)

        if tag in HEADING_TAGS:
            self._render_heading(node)
            return
        if tag == "p":
            if len(node) == 1 and node[0].tag == "img" and not (node.text or "").strip():
                self._render_image(node[0])
            else:
                self._render_paragraph(node)
            return
        if tag in {"ul", "ol"}:
            self._render_list(node)
            return
        if tag == "blockquote":
            self._render_blockquote(node)
            return
        if tag == "pre":
            self._render_pre(node)
            return
        if tag == "table":
            self._render_table(node)
            return
        if tag == "hr":
            if node.get("class") == "__pagebreak__":
                self.document.add_page_break()
            else:
                self._render_hr()
            return
        if tag == "figure":
            for child in node:
                self._render_block(child)
            return
        if tag == "img":
            self._render_image(node)
            return
        if tag in {"div", "section", "article", "nav", "aside"}:
            for child in node:
                self._render_block(child)
            return
        if tag in {"tbody", "thead", "tfoot", "tr", "td", "th"}:
            return

        if getattr(node, "text", None) and node.text.strip():
            paragraph = self.document.add_paragraph()
            self._render_inline_children(node, paragraph)
        for child in node:
            self._render_block(child)
            if child.tail and child.tail.strip():
                paragraph = self.document.add_paragraph()
                self._add_text(paragraph, child.tail)

    def _render_heading(self, node) -> None:
        level = int(node.tag[1])
        paragraph = self.document.add_paragraph(style=f"Heading {level}")
        paragraph.paragraph_format.space_before = Pt(8)
        paragraph.paragraph_format.space_after = Pt(4)
        paragraph.add_run(text_content(node))
        paragraph.runs[0].font.color.rgb = self.title_color

        bookmark = self._bookmark_for(node.get("id") or slugify(text_content(node)))
        if bookmark:
            add_bookmark(paragraph, bookmark, self._next_bookmark_id())

    def _render_paragraph(self, node) -> None:
        paragraph = self.document.add_paragraph()
        paragraph.paragraph_format.space_after = Pt(6)
        paragraph.paragraph_format.space_before = Pt(0)
        self._render_inline_children(node, paragraph)

    def _render_inline_children(self, node, paragraph) -> None:
        if node.text:
            self._add_text(paragraph, node.text)
        for child in node:
            self._render_inline_node(child, paragraph, InlineStyle())
            if child.tail:
                self._add_text(paragraph, child.tail)

    def _render_inline_node(self, node, paragraph, style: InlineStyle) -> None:
        tag = getattr(node, "tag", None)

        if tag == "br":
            paragraph.add_run().add_break()
            return
        if tag == "a":
            href = node.get("href", "")
            label = text_content(node) or href
            if href.startswith("#"):
                bookmark = self._bookmark_for(href)
                if bookmark:
                    self._add_hyperlink(paragraph, label, bookmark, internal=True, style=style)
                else:
                    self._add_text(paragraph, label, style)
            else:
                self._add_hyperlink(paragraph, label, href, internal=False, style=style)
            return
        if tag == "img":
            self._render_image(node, paragraph)
            return

        next_style = InlineStyle(
            bold=style.bold or tag in {"strong", "b"},
            italic=style.italic or tag in {"em", "i"},
            code=style.code or tag == "code",
            strike=style.strike or tag in {"del", "s"},
            underline=style.underline or tag == "u",
            color=style.color,
        )

        if node.text:
            self._add_text(paragraph, node.text, next_style)
        for child in node:
            self._render_inline_node(child, paragraph, next_style)
            if child.tail:
                self._add_text(paragraph, child.tail, style)

    def _render_list(self, node, depth: int = 0) -> None:
        style_name = "List Number" if node.tag == "ol" else "List Bullet"
        for li in node.xpath("./li"):
            paragraph = self.document.add_paragraph(style=style_name)
            paragraph.paragraph_format.left_indent = Inches(0.25 * depth)
            paragraph.paragraph_format.space_after = Pt(0)
            paragraph.paragraph_format.space_before = Pt(0)

            if li.text and li.text.strip():
                self._add_text(paragraph, li.text)

            for child in li:
                if child.tag in {"ul", "ol"}:
                    self._render_list(child, depth + 1)
                else:
                    self._render_inline_node(child, paragraph, InlineStyle())
                if child.tail and child.tail.strip():
                    self._add_text(paragraph, child.tail)

    def _render_blockquote(self, node) -> None:
        for child in node:
            paragraph = self.document.add_paragraph()
            paragraph.paragraph_format.left_indent = Inches(0.3)
            paragraph.paragraph_format.space_before = Pt(2)
            paragraph.paragraph_format.space_after = Pt(2)
            add_paragraph_shading(paragraph, self.theme["quote_fill"])
            self._render_block(child)

    def _render_pre(self, node) -> None:
        paragraph = self.document.add_paragraph()
        paragraph.paragraph_format.left_indent = Inches(0.15)
        paragraph.paragraph_format.right_indent = Inches(0.15)
        paragraph.paragraph_format.space_before = Pt(4)
        paragraph.paragraph_format.space_after = Pt(6)
        add_paragraph_shading(paragraph, self.theme["code_fill"])
        add_paragraph_border(paragraph)
        code_text = text_content(node)
        run = paragraph.add_run(code_text)
        run.font.name = self.code_font
        run.font.size = Pt(9)

    def _render_table(self, node) -> None:
        rows = node.xpath(".//tr")
        if not rows:
            return

        column_count = max(len(r.xpath("./th|./td")) for r in rows)
        table = self.document.add_table(rows=0, cols=column_count)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.style = "Table Grid"

        for row_index, row in enumerate(rows):
            cells = row.xpath("./th|./td")
            docx_row = table.add_row()
            for col_index in range(column_count):
                cell = docx_row.cells[col_index]
                cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
                if col_index < len(cells):
                    if row_index == 0:
                        add_cell_shading(cell, self.theme["table_head_fill"])
                    self._render_table_cell(cell, cells[col_index])
                else:
                    cell.text = ""

    def _render_table_cell(self, cell, node) -> None:
        if node.text and node.text.strip():
            paragraph = cell.paragraphs[0]
            self._render_inline_children(node, paragraph)
        for child in node:
            if child.tag in BLOCK_TAGS:
                if child.tag == "p":
                    paragraph = cell.add_paragraph()
                    self._render_inline_children(child, paragraph)
                else:
                    self._render_block(child)
            else:
                paragraph = cell.paragraphs[0]
                self._render_inline_node(child, paragraph, InlineStyle())
            if child.tail and child.tail.strip():
                paragraph = cell.add_paragraph()
                self._add_text(paragraph, child.tail)

    def _render_image(self, node, paragraph=None) -> None:
        src = node.get("src", "")
        alt = node.get("alt", "Imagen")
        image_bytes, resolved = read_image_bytes(src, self.source_dir)
        target_paragraph = paragraph or self.document.add_paragraph()
        if image_bytes is None:
            self._add_text(target_paragraph, f"[Imagen no encontrada: {src or resolved or alt}]")
            return
        run = target_paragraph.add_run()
        try:
            run.add_picture(image_bytes, width=self.available_width)
        except Exception as exc:
            logger.warning("No se pudo insertar la imagen '%s' (%s): %s", alt, src, exc)
            self._add_text(target_paragraph, f"[No fue posible insertar la imagen: {alt}]")

    def _render_hr(self) -> None:
        paragraph = self.document.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = paragraph.add_run("─" * 24)
        run.font.color.rgb = RGBColor(*self.theme["hr_color"])


def validate_markdown_extension(source_path: Path) -> None:
    if source_path.suffix.lower() not in MARKDOWN_EXTENSIONS:
        allowed = ", ".join(sorted(MARKDOWN_EXTENSIONS))
        raise ValueError(
            f"'{source_path.name}' no tiene una extensión Markdown reconocida ({allowed}). "
            "Verifica que el archivo sea realmente Markdown antes de convertir."
        )


def validate_source_extension(source_path: Path) -> None:
    """Valida que la extensión sea Markdown o HTML."""
    if source_path.suffix.lower() not in SUPPORTED_EXTENSIONS:
        allowed = ", ".join(sorted(SUPPORTED_EXTENSIONS))
        raise ValueError(
            f"'{source_path.name}' no tiene una extensión soportada ({allowed}). "
            "Usa archivos .md, .markdown, .html o .htm."
        )


def detect_inline_tables(markdown_text: str) -> list[int]:
    """Devuelve los números de línea (1-based) donde una tabla empieza sin línea en blanco antes.

    Markdown requiere una línea en blanco entre un párrafo y una tabla; si no, el motor
    trata la tabla como continuación del párrafo y no la renderiza como tabla. Esta función
    es solo diagnóstico: NO modifica el texto, solo avisa al usuario para que lo arregle.
    """
    issues: list[int] = []
    lines = markdown_text.split("\n")
    for i, line in enumerate(lines):
        if not line.lstrip().startswith("|"):
            continue
        if i == 0:
            continue
        prev = lines[i - 1]
        if prev.strip() == "":
            continue
        if prev.lstrip().startswith("|"):
            continue
        issues.append(i + 1)
    return issues


def warn_inline_tables(source_path: Path, markdown_text: str) -> None:
    """Emite una advertencia por cada tabla 'inline' encontrada en el .md."""
    bad_lines = detect_inline_tables(markdown_text)
    if not bad_lines:
        return
    plural = "s" if len(bad_lines) > 1 else ""
    logger.warning(
        "⚠ %s contiene %d tabla%s sin línea en blanco antes "
        "(línea%s %s). Esas tablas se renderizarán como texto plano, no como tabla. "
        "Añade una línea en blanco justo antes de cada '| ... |' para corregirlo.",
        source_path.name,
        len(bad_lines),
        plural,
        plural,
        ", ".join(str(n) for n in bad_lines),
    )


def _source_to_html_body(source_path: Path) -> str:
    """Lee un archivo .md o .html y devuelve el HTML del body listo para combinar."""
    text = source_path.read_text(encoding="utf-8")
    if source_path.suffix.lower() in HTML_EXTENSIONS:
        body_match = re.search(r"<body[^>]*>(.*)</body>", text, re.IGNORECASE | re.DOTALL)
        return body_match.group(1) if body_match else text
    return markdown_to_html(text)


def convert_source_to_docx(
    source_path: Path,
    output_path: Path,
    theme: dict | None = None,
) -> Document:
    """Convierte un .md/.html a un Document de python-docx (sin guardar)."""
    text = source_path.read_text(encoding="utf-8")
    if source_path.suffix.lower() in MARKDOWN_EXTENSIONS:
        warn_inline_tables(source_path, text)
    converter = MarkdownToDocxConverter(source_path, output_path, theme=theme)
    if source_path.suffix.lower() in HTML_EXTENSIONS:
        return converter.convert_html(text)
    return converter.convert_markdown(text)


def convert_source_file(
    source_path: Path,
    output_path: Path | None = None,
    theme_path: Path | None = None,
) -> Path:
    """Convierte un .md o .html a un .docx y lo guarda en disco."""
    source_path = Path(source_path).resolve()
    if not source_path.exists():
        raise FileNotFoundError(f"No existe el archivo de entrada: {source_path}")
    validate_source_extension(source_path)

    if output_path is None:
        output_path = source_path.with_suffix(".docx")
    output_path = Path(output_path).resolve()

    theme = load_theme(theme_path)
    document = convert_source_to_docx(source_path, output_path, theme=theme)
    document.save(output_path)
    logger.info("Documento creado: %s", output_path)
    return output_path


def convert_markdown_file(
    source_path: Path,
    output_path: Path | None = None,
    theme_path: Path | None = None,
) -> Path:
    """Compatibilidad hacia atrás: alias de convert_source_file()."""
    return convert_source_file(source_path, output_path, theme_path)


def _expand_patterns(patterns: list[str]) -> list[Path]:
    """Resuelve una mezcla de rutas literales y patrones glob a una lista única de archivos."""
    seen: set[Path] = set()
    resolved: list[Path] = []
    for pattern in patterns:
        pattern_path = Path(pattern)
        if pattern_path.exists() and pattern_path.is_file():
            target = pattern_path.resolve()
        else:
            matches = sorted(Path(p).resolve() for p in glob.glob(pattern))
            target = None
            for m in matches:
                if m.is_file() and m not in seen:
                    seen.add(m)
                    resolved.append(m)
            continue
        if target not in seen:
            seen.add(target)
            resolved.append(target)
    return resolved


def convert_many(
    patterns: list[str],
    output_dir: Path | None = None,
    theme_path: Path | None = None,
) -> list[Path]:
    """Convierte múltiples archivos .md/.html (literales o glob) en .docx separados."""
    source_paths = _expand_patterns(patterns)
    if not source_paths:
        raise FileNotFoundError(f"Ningún archivo coincide con los patrones: {patterns}")

    results: list[Path] = []
    for source_path in source_paths:
        try:
            validate_source_extension(source_path)
        except ValueError as exc:
            logger.warning("Omitiendo %s: %s", source_path, exc)
            continue

        output_path = None
        if output_dir is not None:
            output_path = Path(output_dir) / source_path.with_suffix(".docx").name

        result = convert_source_file(source_path, output_path, theme_path=theme_path)
        results.append(result)

    return results


def convert_merged_docx(
    source_paths: list[Path],
    output_path: Path,
    theme: dict | None = None,
) -> Path:
    """Une varios .md/.html en un único .docx, en el orden recibido.

    Cada archivo ocupa su propio bloque con salto de página entre ellos. Los
    marcadores ``<!-- pagebreak -->`` dentro de un mismo archivo también se
    respetan (esa lógica ya vive en ``MarkdownToDocxConverter.convert_markdown``).
    """
    if not source_paths:
        raise ValueError("convert_merged_docx requiere al menos un archivo de entrada.")

    output_path = Path(output_path).resolve()
    output_path.parent.mkdir(parents=True, exist_ok=True)

    if theme is None:
        theme = dict(DEFAULT_THEME)

    first = Path(source_paths[0]).resolve()
    document = Document()
    configure_document(document, theme)

    for index, raw in enumerate(source_paths):
        path = Path(raw).resolve()
        if not path.exists():
            raise FileNotFoundError(f"No existe el archivo de entrada: {path}")
        if path.suffix.lower() not in SUPPORTED_EXTENSIONS:
            allowed = ", ".join(sorted(SUPPORTED_EXTENSIONS))
            raise ValueError(f"'{path.name}' no es una entrada soportada ({allowed}).")

        if path.suffix.lower() in MARKDOWN_EXTENSIONS:
            warn_inline_tables(path, path.read_text(encoding="utf-8"))

        # Construimos un converter "temporal" que reusa el mismo Document
        # para que el primer archivo no pierda el formato base.
        if index == 0:
            converter = MarkdownToDocxConverter(path, output_path, theme=theme)
            converter.document = document  # reutilizamos el doc recién creado
            if path.suffix.lower() in HTML_EXTENSIONS:
                converter.convert_html(path.read_text(encoding="utf-8"))
            else:
                converter.convert_markdown(path.read_text(encoding="utf-8"))
        else:
            # Para los siguientes, partimos de cero y luego movemos todos los
            # elementos del body al documento principal.
            sub_converter = MarkdownToDocxConverter(path, output_path, theme=theme)
            if path.suffix.lower() in HTML_EXTENSIONS:
                sub_converter.convert_html(path.read_text(encoding="utf-8"))
            else:
                sub_converter.convert_markdown(path.read_text(encoding="utf-8"))

            # Salto de página entre archivos
            document.add_page_break()
            _merge_documents(document, sub_converter.document)

    document.save(output_path)
    logger.info("Documento unificado creado: %s", output_path)
    return output_path


def _merge_documents(target: Document, source: Document) -> None:
    """Copia los elementos de body de ``source`` al final de ``target`` (sin la sectPr final)."""
    from copy import deepcopy

    src_body = source.element.body
    target_body = target.element.body
    # Conservamos el último elemento de ``target_body`` (la sectPr del documento)
    # y añadimos todos los hijos de ``src_body`` antes de él.
    sectPr = target_body.find(qn("w:sectPr"))
    for child in list(src_body):
        if child.tag == qn("w:sectPr"):
            continue
        target_body.insert(list(target_body).index(sectPr) if sectPr is not None else len(target_body), deepcopy(child))


def convert_merged_from_patterns(
    patterns: list[str],
    output_path: Path,
    theme_path: Path | None = None,
) -> Path:
    """Atajo: resuelve patrones/glob y une todos los .md/.html en un solo .docx."""
    source_paths = _expand_patterns(patterns)
    if not source_paths:
        raise FileNotFoundError(f"Ningún archivo coincide con los patrones: {patterns}")
    for path in source_paths:
        validate_source_extension(path)
    theme = load_theme(theme_path)
    return convert_merged_docx(source_paths, output_path, theme=theme)


def watch_and_convert(source_path: Path, output_path: Path | None, theme_path: Path | None) -> None:
    """Vigila un archivo y reconvierte automáticamente cada vez que cambia."""
    source_path = Path(source_path).resolve()
    last_mtime: float | None = None
    logger.info("Modo watch activo sobre %s (Ctrl+C para detener)", source_path)
    try:
        while True:
            if not source_path.exists():
                time.sleep(1)
                continue
            mtime = source_path.stat().st_mtime
            if mtime != last_mtime:
                last_mtime = mtime
                try:
                    convert_source_file(source_path, output_path, theme_path=theme_path)
                except Exception as exc:
                    logger.error("Falló la reconversión: %s", exc)
            time.sleep(1)
    except KeyboardInterrupt:
        logger.info("Modo watch detenido.")


def build_parser() -> argparse.ArgumentParser:
    parser = argparse.ArgumentParser(
        description="Convierte uno o más archivos Markdown/HTML a Word (.docx) con formato conservado."
    )
    parser.add_argument(
        "input", nargs="+",
        help="Ruta(s) o patrón(es) glob de archivo(s) .md/.markdown/.html/.htm de entrada",
    )
    parser.add_argument(
        "-o", "--output",
        help="Ruta de salida .docx (un solo archivo) o carpeta de salida (batch)",
    )
    parser.add_argument(
        "--theme",
        help=(
            "Nombre de tema integrado (professional, minimal, empresarial, "
            "academico, economico, explicativo) o ruta a un archivo "
            ".json/.toml/.yaml con overrides"
        ),
    )
    parser.add_argument(
        "--merge", action="store_true",
        help="Une todos los archivos de entrada en un único .docx (en orden)",
    )
    parser.add_argument(
        "--watch", action="store_true",
        help="Vigila el archivo de entrada y reconvierte al detectar cambios (solo un archivo)",
    )
    return parser


def main(argv: list[str] | None = None) -> int:
    parser = build_parser()
    args = parser.parse_args(argv)

    theme_path = Path(args.theme) if args.theme else None

    try:
        if args.watch:
            if len(args.input) != 1:
                print("Error: --watch solo admite un único archivo de entrada.")
                return 1
            watch_and_convert(Path(args.input[0]), Path(args.output) if args.output else None, theme_path)
            return 0

        if args.merge:
            output = Path(args.output) if args.output else Path("merged.docx").resolve()
            result = convert_merged_from_patterns(args.input, output, theme_path=theme_path)
            print(f"Documento unificado creado: {result}")
            return 0

        if len(args.input) == 1 and Path(args.input[0]).exists() and Path(args.input[0]).is_file():
            output_path = Path(args.output) if args.output else None
            result = convert_source_file(Path(args.input[0]), output_path, theme_path=theme_path)
            print(f"Documento creado: {result}")
            return 0

        output_dir = Path(args.output) if args.output else None
        results = convert_many(args.input, output_dir=output_dir, theme_path=theme_path)
        for result in results:
            print(f"Documento creado: {result}")
        print(f"Total: {len(results)} documento(s) generado(s).")
        return 0
    except (FileNotFoundError, ValueError, RuntimeError) as exc:
        print(f"Error: {exc}")
        return 1


if __name__ == "__main__":
    raise SystemExit(main())
