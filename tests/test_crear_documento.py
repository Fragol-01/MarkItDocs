"""Self-check de humo: valida que convert_markdown_file produce un .docx real."""

import json
import sys
import tempfile
from pathlib import Path

sys.path.insert(0, str(Path(__file__).resolve().parent.parent))

from crear_documento import (  # noqa: E402
    convert_many,
    convert_markdown_file,
    convert_merged_from_patterns,
    convert_source_file,
    validate_markdown_extension,
    validate_source_extension,
)
from docx import Document  # noqa: E402

SAMPLE_HTML = """<!doctype html>
<html>
<head><title>Documento HTML</title></head>
<body>
<h1>Título HTML</h1>
<p>Párrafo con <strong>negrita</strong> y <em>cursiva</em>.</p>
<table><tr><th>A</th><th>B</th></tr><tr><td>1</td><td>2</td></tr></table>
</body>
</html>
"""

SAMPLE_MARKDOWN = """# Documento de prueba

## Sección con tabla

| Col A | Col B |
|---|---|
| 1 | 2 |

- item uno
- item dos

```python
print("hola")
```

> Una cita breve.

[Enlace externo](https://example.com)
"""


def test_convert_produces_nonempty_docx():
    with tempfile.TemporaryDirectory() as tmp_dir:
        md_path = Path(tmp_dir) / "sample.md"
        md_path.write_text(SAMPLE_MARKDOWN, encoding="utf-8")

        output = convert_markdown_file(md_path)

        assert output.exists()
        assert output.stat().st_size > 1000, "El .docx generado es sospechosamente pequeño"
        assert output.suffix == ".docx"


def test_convert_with_missing_local_image():
    markdown_with_broken_image = SAMPLE_MARKDOWN + "\n![alt](no_existe.png)\n"
    with tempfile.TemporaryDirectory() as tmp_dir:
        md_path = Path(tmp_dir) / "sample.md"
        md_path.write_text(markdown_with_broken_image, encoding="utf-8")

        output = convert_markdown_file(md_path)

        assert output.exists(), "Debe generar el .docx aunque falte una imagen local"


def test_rejects_non_markdown_extension():
    with tempfile.TemporaryDirectory() as tmp_dir:
        txt_path = Path(tmp_dir) / "fake.txt"
        txt_path.write_text("no soy markdown", encoding="utf-8")

        try:
            validate_markdown_extension(txt_path)
        except ValueError:
            pass
        else:
            raise AssertionError("Se esperaba ValueError para extensión .txt")


def test_pagebreak_inserts_real_page_break():
    markdown_with_break = "# Pagina 1\ntexto\n\n\\pagebreak\n\n# Pagina 2\ntexto\n"
    with tempfile.TemporaryDirectory() as tmp_dir:
        md_path = Path(tmp_dir) / "pb.md"
        md_path.write_text(markdown_with_break, encoding="utf-8")

        output = convert_markdown_file(md_path)
        document = Document(output)

        xml = document.element.xml
        assert 'type="page"' in xml, "Debe insertar un salto de página real en el XML"


def test_custom_theme_overrides_defaults():
    with tempfile.TemporaryDirectory() as tmp_dir:
        md_path = Path(tmp_dir) / "sample.md"
        md_path.write_text(SAMPLE_MARKDOWN, encoding="utf-8")

        theme_path = Path(tmp_dir) / "theme.json"
        theme_path.write_text(json.dumps({"title_color": [200, 30, 30], "body_font": "Georgia"}), encoding="utf-8")

        output = convert_markdown_file(md_path, theme_path=theme_path)
        document = Document(output)

        assert document.styles["Normal"].font.name == "Georgia"
        heading_color = document.paragraphs[0].runs[0].font.color.rgb
        assert str(heading_color) == "C81E1E"


def test_convert_many_batch():
    with tempfile.TemporaryDirectory() as tmp_dir:
        (Path(tmp_dir) / "a.md").write_text("# Doc A\ntexto\n", encoding="utf-8")
        (Path(tmp_dir) / "b.md").write_text("# Doc B\ntexto\n", encoding="utf-8")

        results = convert_many([str(Path(tmp_dir) / "*.md")])

        assert len(results) == 2
        assert all(r.exists() for r in results)


def test_html_to_docx():
    with tempfile.TemporaryDirectory() as tmp_dir:
        html_path = Path(tmp_dir) / "sample.html"
        html_path.write_text(SAMPLE_HTML, encoding="utf-8")

        output = convert_source_file(html_path)

        assert output.exists() and output.suffix == ".docx"
        document = Document(output)
        texts = [p.text for p in document.paragraphs]
        assert any("Título HTML" in t for t in texts), "El h1 del HTML debe estar en el .docx"


def test_html_extension_accepted_md_only_validator_rejects():
    with tempfile.TemporaryDirectory() as tmp_dir:
        html_path = Path(tmp_dir) / "x.html"
        html_path.write_text("<p>hola</p>", encoding="utf-8")
        validate_source_extension(html_path)  # no debe lanzar
        try:
            validate_markdown_extension(html_path)
        except ValueError:
            pass
        else:
            raise AssertionError("validate_markdown_extension debe rechazar .html")


def test_merged_docx_preserves_order_and_page_breaks():
    with tempfile.TemporaryDirectory() as tmp_dir:
        (Path(tmp_dir) / "01_intro.md").write_text("# Introducción\ncontenido uno\n", encoding="utf-8")
        (Path(tmp_dir) / "02_medio.html").write_text("<h1>Medio</h1><p>contenido dos</p>", encoding="utf-8")
        (Path(tmp_dir) / "03_fin.md").write_text("# Final\ncontenido tres\n", encoding="utf-8")

        output = Path(tmp_dir) / "unido.docx"
        result = convert_merged_from_patterns(
            [
                str(Path(tmp_dir) / "01_intro.md"),
                str(Path(tmp_dir) / "02_medio.html"),
                str(Path(tmp_dir) / "03_fin.md"),
            ],
            output,
        )

        assert result.exists()
        document = Document(result)
        texts = [p.text for p in document.paragraphs if p.text.strip()]
        idx_intro = next(i for i, t in enumerate(texts) if "Introducción" in t)
        idx_medio = next(i for i, t in enumerate(texts) if "Medio" in t)
        idx_fin = next(i for i, t in enumerate(texts) if "Final" in t)
        assert idx_intro < idx_medio < idx_fin, "El orden de los archivos debe conservarse"

        xml = document.element.xml
        assert xml.count('type="page"') >= 2, "Debe haber salto de página entre cada archivo"


def test_load_theme_by_builtin_name():
    theme = __import__("crear_documento").load_theme("academico")
    assert theme["body_font"] == "Times New Roman"
    assert theme["title_color"] == [26, 26, 46]
    assert "#" not in theme["table_head_fill"]

    try:
        __import__("crear_documento").load_theme("tema_inexistente")
    except ValueError:
        pass
    else:
        raise AssertionError("Se esperaba ValueError para tema integrado inexistente")


def test_table_without_preceding_blank_line_renders():
    md = "Texto pegado a la tabla:\n| A | B |\n| :--- | :--- |\n| 1 | 2 |\n"
    with tempfile.TemporaryDirectory() as tmp_dir:
        md_path = Path(tmp_dir) / "t.md"
        md_path.write_text(md, encoding="utf-8")
        output = convert_markdown_file(md_path)
        document = Document(output)
        assert len(document.tables) == 1, "La tabla sin línea en blanco previa debe renderizarse"


def test_yaml_theme_loads():
    with tempfile.TemporaryDirectory() as tmp_dir:
        md_path = Path(tmp_dir) / "sample.md"
        md_path.write_text(SAMPLE_MARKDOWN, encoding="utf-8")

        theme_path = Path(tmp_dir) / "theme.yaml"
        theme_path.write_text("body_font: Verdana\ntitle_color: [10, 100, 200]\n", encoding="utf-8")

        output = convert_markdown_file(md_path, theme_path=theme_path)
        document = Document(output)
        assert document.styles["Normal"].font.name == "Verdana"


if __name__ == "__main__":
    test_convert_produces_nonempty_docx()
    print("OK: test_convert_produces_nonempty_docx")
    test_convert_with_missing_local_image()
    print("OK: test_convert_with_missing_local_image")
    test_rejects_non_markdown_extension()
    print("OK: test_rejects_non_markdown_extension")
    test_pagebreak_inserts_real_page_break()
    print("OK: test_pagebreak_inserts_real_page_break")
    test_custom_theme_overrides_defaults()
    print("OK: test_custom_theme_overrides_defaults")
    test_convert_many_batch()
    print("OK: test_convert_many_batch")
    test_html_to_docx()
    print("OK: test_html_to_docx")
    test_html_extension_accepted_md_only_validator_rejects()
    print("OK: test_html_extension_accepted_md_only_validator_rejects")
    test_merged_docx_preserves_order_and_page_breaks()
    print("OK: test_merged_docx_preserves_order_and_page_breaks")
    test_yaml_theme_loads()
    print("OK: test_yaml_theme_loads")
    test_load_theme_by_builtin_name()
    print("OK: test_load_theme_by_builtin_name")
    test_table_without_preceding_blank_line_renders()
    print("OK: test_table_without_preceding_blank_line_renders")
    print("Todos los self-checks pasaron.")
